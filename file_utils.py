import logging
import fitz
from langchain_core.tools import ToolException
from datetime import datetime
from docx import Document
import pdfkit
import re
import os

def extract_text_from_pdf(file):
    """Extracts text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        logging.info("Text extracted from PDF successfully.")
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        raise ToolException(f"Error processing PDF: {e}")
    
def save_as_pdf(content: str, job_title: str, company_name: str) -> str:
    """Saves content as a PDF file."""
    try:
        # Clean filename to avoid special characters
        safe_title = re.sub(r'[\\/*?:"<>|]', "", job_title)
        safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create a temporary HTML file
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Cover Letter - {safe_title} at {safe_company}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 1in; line-height: 1.5; }}
                .header {{ margin-bottom: 20px; }}
                .content {{ margin-bottom: 20px; white-space: pre-wrap; }}
                .signature {{ margin-top: 40px; }}
            </style>
        </head>
        <body>
            <div class="content">{content}</div>
        </body>
        </html>
        """
        
        # Create directory if it doesn't exist
        output_dir = os.path.join(os.getcwd(), "cover_letters")
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate file paths
        html_path = os.path.join(output_dir, f"cover_letter_{safe_title}_{safe_company}_{timestamp}.html")
        pdf_path = os.path.join(output_dir, f"cover_letter_{safe_title}_{safe_company}_{timestamp}.pdf")
        
        # Write HTML file
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        # Convert to PDF
        try:
            pdfkit.from_file(html_path, pdf_path)
            os.remove(html_path)  # Clean up HTML file
            logging.info(f"Cover letter saved as PDF: {pdf_path}")
            return pdf_path
        except Exception as e:
            logging.error(f"Error converting to PDF, falling back to HTML: {e}")
            return html_path
            
    except Exception as e:
        logging.error(f"Error saving cover letter as PDF: {e}")
        raise ToolException(f"Failed to save cover letter as PDF: {e}")

def save_as_docx(content: str, job_title: str, company_name: str) -> str:
    """Saves content as a DOCX file."""
    try:
        # Clean filename to avoid special characters
        safe_title = re.sub(r'[\\/*?:"<>|]', "", job_title)
        safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create directory if it doesn't exist
        output_dir = os.path.join(os.getcwd(), "cover_letters")
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate file path
        docx_path = os.path.join(output_dir, f"cover_letter_{safe_title}_{safe_company}_{timestamp}.docx")
        
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading(f"Cover Letter - {job_title} at {company_name}", 1)
        
        # Add content (split by paragraphs)
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():  # Skip empty lines
                doc.add_paragraph(para)
        
        # Save document
        doc.save(docx_path)
        logging.info(f"Cover letter saved as DOCX: {docx_path}")
        return docx_path
        
    except Exception as e:
        logging.error(f"Error saving cover letter as DOCX: {e}")
        raise ToolException(f"Failed to save cover letter as DOCX: {e}")
